#!/usr/bin/env python3

#from abbreviations import schwartz_hearst
from pydoc import doc
from schwartz_hearst import extract_abbreviation_definition_pairs
import json
import sys, getopt
from venv import create
import docx2txt
import re
from flask import jsonify
import xlsxwriter
import PyPDF2 as pdf2
import csv
from docx import Document
from collections import OrderedDict

#stores dictionary of acronyms and abbreviations
dict_from_csv = {}

#opens given file path and finds full abbreviations of acronyms
def findAbbrev(inputfile):
    # TO-DO: Remove parenthesis from the outside
    # doc_text = inputfile
    doc_text=docx2txt.process(inputfile)
    openingParenthesisStack = []
    # iterates through index and character of enumerated text
    for i, char in enumerate(doc_text):
        # checks for open and closing parenthesis and appends into stack
        if char == '(':
            openingParenthesisStack.append((i, char))
        if char == ')':
            if openingParenthesisStack:
                # the last item in stack is '(' pop the last item in stack
                if openingParenthesisStack[len(openingParenthesisStack) - 1][1] == '(':
                    openingParenthesisStack.pop()
            # Condition: '()))'
            else:                
                doc_text = doc_text[:i] + doc_text[i+1:]
    if openingParenthesisStack:
        # keep on replacing the '(' from text with the index and character obtained from the pop item from stack
        while openingParenthesisStack:
            i, char = openingParenthesisStack.pop()
            doc_text = doc_text[:i] + doc_text[i+1:]
    dict_from_csv = extract_abbreviation_definition_pairs(doc_text=doc_text)
    # pairs = ''
    # global dict_from_csv
    # dict_from_csv = OrderedDict(sorted(pairs.items(), key=lambda t: t[0]))

    # dict_from_csv = OrderedDict(sorted(pairs.items(), key=lambda t: t[0]))

    return dict_from_csv

#revert back to number short forms
def revertToNumber(word):
    counter = 0
    listOfWords = 0
    if(word[0] == word[1]):
        for char in word:
            if(word[0] == char):
                counter += 1
            else:
                break
        
        word = word[:counter] + str(counter) + word[counter:]
        ch = word[counter-1]
        listOfWords = word.split(ch, counter-1)
        if len(listOfWords) > 0:
            word = listOfWords[len(listOfWords)-1]

    return word

#return dict_from_csv
def getDictFromCSV(inputfile):
    data = findAbbrev(inputfile)
    return data

#opens document and returns acronyms found
def findAcronyms(inputfile):
    outputfile = "list"

    rx = r"\b[A-Z](?=([&.]?))(?:\1[A-Z])+\b"
    doc_text=docx2txt.process(inputfile)
    list = [x.group() for x in re.finditer(rx, doc_text)]

    # remove repeated acronyms from list
    aclist = []
    for i in list:
        if i not in aclist:
           aclist.append(str(i))
    
    #obtains list of acronyms not found in shwartz-hearst algorithm
    notlist = []
    global dict_from_csv
    dict_from_csv_copy = dict_from_csv.copy()
    for item in aclist:
        if item in dict_from_csv_copy:
            pass
        else:
            dict_from_csv_copy.update({item: " "})
    dict_from_csv_copy = OrderedDict(sorted(dict_from_csv_copy.items(), key=lambda t: t[0]))
    # print(dict_from_csv_copy)    

    documentObj = Document()

    table = documentObj.add_table(rows=1, cols=2)

    check = []
    
    for k,v in dict_from_csv_copy.items():
        row_cells = table.add_row().cells
        row_cells[0].text = k
        row_cells[1].text = v

    documentObj.add_page_break()

    documentObj.save("static/client/docx/" + outputfile + '.docx')

    return check

#iterates through given document and creates an excel workbook
#with given outputfile name
def createXLSX(outfilename, acronymdata):
    filename = (outfilename).split(".docx")[0] + ".xlsx"

    # sorts dictionary in alphabetical order
    acronymdata = OrderedDict(sorted(acronymdata.items(), key=lambda t: t[0]))
    # creates xlsx object
    workbook = xlsxwriter.Workbook("static/client/xlsx/" + filename)
    worksheet = workbook.add_worksheet()

    row = 0
    for k,v in acronymdata.items():
        worksheet.write(row, 0, revertToNumber(k))
        worksheet.write(row, 1, v)
        
        row += 1
    workbook.close()

#iterates through given document and creates a word docx
#with given outputfile name
def createDoc(outfilename, acronymdata):
    # sorts dictionary in alphabetical order
    acronymdata = OrderedDict(sorted(acronymdata.items(), key=lambda t: t[0]))

    # creates document object with one row and two columns
    documentObj = Document()
    table = documentObj.add_table(rows=1, cols=2)
    # iterates through global dict_from_csv to create table
    for k,v in acronymdata.items():
        row_cells = table.add_row().cells
        row_cells[0].text = revertToNumber(k)
        row_cells[1].text = v

    documentObj.add_page_break()

    documentObj.save("static/client/docx/" + outfilename)



